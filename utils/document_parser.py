import os
import re
import docx
import PyPDF2
from typing import Dict, List, Tuple

def extract_text_from_document(file_path: str) -> str:
    """
    Extract text from various document formats (PDF, DOCX, TXT)
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Extracted text content as a string
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text = ""
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    return clean_extracted_text(text)

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    text = ""
    
    try:
        doc = docx.Document(file_path)
        
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    return clean_extracted_text(text)

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    except UnicodeDecodeError:
        # Try with a different encoding if UTF-8 fails
        with open(file_path, 'r', encoding='latin-1') as file:
            text = file.read()
    except Exception as e:
        raise Exception(f"Error extracting text from TXT: {str(e)}")
    
    return clean_extracted_text(text)

def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Fix common OCR/extraction issues
    text = text.replace('|', 'I').replace('1', 'l')
    
    # Remove headers/footers that might be repeated on pages
    # This is a simplified approach; may need refinement for complex documents
    lines = text.split('\n')
    cleaned_lines = []
    
    for i, line in enumerate(lines):
        # Skip potential headers/footers (very short lines that appear multiple times)
        if len(line.strip()) < 50 and lines.count(line) > 2:
            continue
        cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def identify_clauses_regex(text: str) -> Dict[str, str]:
    """
    Attempt to identify contract clauses using regex patterns
    This is a fallback method if the LLM-based extraction fails
    
    Args:
        text: The contract text
        
    Returns:
        Dictionary of clause titles and their content
    """
    # Common patterns for clause numbering
    patterns = [
        r'(\d+\.\s*[A-Z][^\.]+)(?:\.|:)(.*?)(?=\d+\.\s*[A-Z][^\.]+(?:\.|:)|$)', # 1. Title: Content
        r'([A-Z][A-Z\s]+)(?:\.|:)(.*?)(?=[A-Z][A-Z\s]+(?:\.|:)|$)',  # ALL CAPS TITLE: Content
        r'((?:Article|Section|Clause)\s+\d+[^\.]+)(?:\.|:)(.*?)(?=(?:Article|Section|Clause)\s+\d+[^\.]+(?:\.|:)|$)'  # Article 1 - Title: Content
    ]
    
    clauses = {}
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.DOTALL)
        if matches:
            for title, content in matches:
                clean_title = title.strip()
                clean_content = content.strip()
                if clean_title and clean_content:
                    clauses[clean_title] = clean_content
    
    # If we couldn't find clauses with regex, use a simple paragraph-based approach
    if not clauses:
        paragraphs = re.split(r'\n\s*\n', text)
        for i, para in enumerate(paragraphs):
            if len(para.strip()) > 50:  # Only include substantive paragraphs
                clauses[f"Paragraph {i+1}"] = para.strip()
    
    return clauses